from __future__ import annotations
import json
import operator
import os
from pathlib import Path
from typing import Annotated, List, Literal, Optional, TypedDict,ClassVar
from dotenv import load_dotenv
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_groq import ChatGroq
from langgraph.graph import END, START, StateGraph
from langgraph.types import Send
from pydantic import BaseModel, Field, ValidationError,field_validator
from docx import Document
load_dotenv()

# ─────────────────────────────────────────────
# Pydantic Models
# ─────────────────────────────────────────────

class Task(BaseModel):
    id: int
    title: str
    goal: str = Field(
        ...,
        description="One sentence describing what the reader should be able to do/understand after this section.",
    )
    bullets: List[str] = Field(
        ...,
        min_length=2,
        max_length=5,
        description="3–5 concrete, non-overlapping subpoints to cover in this section.",
    )
    target_words: int = Field(
        ...,
        description="Target word count for this section (120–550).",
    )
    tags: List[str] = Field(default_factory=list)
    requires_research: bool = False
    requires_citations: bool = False
    requires_code: bool = False


class Plan(BaseModel):
    blog_title: str
    audience: str
    tone: str
    blog_kind: str = "explainer"
    constraints: List[str] = Field(default_factory=list)
    tasks: List[Task]

    VALID_KINDS: ClassVar[set] = {"explainer", "tutorial", "news_roundup", "comparison", "system_design"}

    @field_validator("blog_kind", mode="before")
    @classmethod
    def fix_blog_kind(cls, v):
        if v not in cls.VALID_KINDS:
            return "explainer"
        return v


class RouterDecision(BaseModel):
    needs_research: bool
    mode: Literal["closed_book", "hybrid", "open_book"]
    queries: List[str] = Field(default_factory=list, max_length=5)


class EvidenceItem(BaseModel):
    title: str
    url: str
    published_at: Optional[str] = None
    snippet: Optional[str] = None
    source: Optional[str] = None


class EvidencePack(BaseModel):
    evidence: List[EvidenceItem] = Field(default_factory=list)


# ─────────────────────────────────────────────
# Graph State
# ─────────────────────────────────────────────

class State(TypedDict):
    topic: str
    mode: str
    needs_research: bool
    queries: List[str]
    evidence: List[EvidenceItem]
    plan: Plan
    sections: Annotated[List[tuple[int, str]], operator.add]
    final: str
    docx_path: str


# ─────────────────────────────────────────────
# LLM Clients
# ─────────────────────────────────────────────

GROQ_API_KEY = os.getenv("GROQ_API")  # make sure your .env uses GROQ_API_KEY

router_llm = ChatGroq(model="llama-3.1-8b-instant", api_key=GROQ_API_KEY)
planner_llm = ChatGroq(model="llama-3.3-70b-versatile", api_key=GROQ_API_KEY, temperature=0)
writer_llm = ChatGroq(model="llama-3.3-70b-versatile", api_key=GROQ_API_KEY, temperature=0.4, max_tokens=600)


# ─────────────────────────────────────────────
# Helper: Tavily Search  ← was missing before
# ─────────────────────────────────────────────

def _tavily_search(query: str, max_results: int = 3) -> List[dict]:
    """Run a Tavily web search and return a list of result dicts."""
    tool = TavilySearchResults(max_results=max_results)
    results = tool.invoke({"query": query})
    # Tavily returns list of dicts: {url, title, content, score, ...}
    return results if isinstance(results, list) else []


# ─────────────────────────────────────────────
# Node: Router
# ─────────────────────────────────────────────

ROUTER_SYSTEM = """
You are a routing module for a technical blog planner.

Return a RouterDecision object.

Rules:

1. needs_research MUST be a boolean.

2. mode MUST be exactly one of:
- closed_book
- hybrid
- open_book

3. queries MUST contain at most 5 search queries.

4. Each query must be a COMPLETE Google search query.

GOOD:
- Open source LLMs 2026 comparison
- Best open source LLMs in 2026
- Open source LLM benchmarks 2026

BAD:
- AI
- blog
- 2026
- LLM

5. If mode == closed_book:
queries MUST be []

6. Never generate more than 5 queries.

Return only the RouterDecision object.
"""


def router_node(state: State) -> dict:
    topic = state["topic"]
    
    response = router_llm.invoke([
        SystemMessage(content=ROUTER_SYSTEM + """
Return ONLY valid JSON like this:
{
  "needs_research": true,
  "mode": "open_book",
  "queries": ["query 1", "query 2"]
}
No markdown. No explanation. Just JSON.
"""),
        HumanMessage(content=f"topic: {topic}"),
    ])

    raw = response.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    data = json.loads(raw)
    decision = RouterDecision(**data)

    return {
        "needs_research": decision.needs_research,
        "mode": decision.mode,
        "queries": decision.queries,
    }

def route_next(state: State) -> str:
    return "research" if state["needs_research"] else "orchestrator"


# ─────────────────────────────────────────────
# Node: Research
# ─────────────────────────────────────────────

def research_node(state: State) -> dict:
    queries = (state.get("queries") or [])[:3]
    max_results = 3

    raw_results: List[dict] = []
    for q in queries:
        raw_results.extend(_tavily_search(q, max_results=max_results))

    if not raw_results:
        return {"evidence": []}

    seen_urls: set[str] = set()
    evidence: List[EvidenceItem] = []

    for r in raw_results:
        url = (r.get("url") or "").strip()
        if not url or url in seen_urls:
            continue
        seen_urls.add(url)

        evidence.append(
            EvidenceItem(
                title=(r.get("title") or "").strip(),
                url=url,
                snippet=(r.get("content") or "")[:300],   # ← Tavily uses "content", not "snippet"
                published_at=r.get("published_date"),      # ← Tavily key is "published_date"
                source=r.get("source"),
            )
        )

    return {"evidence": evidence}


# ─────────────────────────────────────────────
# Node: Orchestrator (Planner)
# ─────────────────────────────────────────────

ORCH_SYSTEM = """
You are a senior technical blog planner.

Return ONLY valid JSON.

Do NOT wrap the JSON inside markdown.
Do NOT explain anything.
Do NOT use code fences.

JSON schema:

{
  "blog_title": "...",
  "audience": "...",
  "tone": "...",
  "blog_kind": "...",
  "constraints": ["...", "..."],
  "tasks": [
    {
      "id": 1,
      "title": "...",
      "goal": "...",
      "bullets": [
        "...",
        "...",
        "..."
      ],
      "target_words": 200,
      "tags": [],
      "requires_research": false,
      "requires_citations": false,
      "requires_code": false
    }
  ]
}

Rules:
- Create 5-7 tasks.
- Every task MUST have all fields.
- bullets must contain at least 3 items.
- constraints must be a JSON array.
- tasks must be a JSON array.
- Output ONLY JSON.
Return ONLY raw JSON.

Never stringify arrays.

Correct:
"constraints": ["AI","LLM"]

Incorrect:
"constraints": "[\"AI\",\"LLM\"]"

Correct:
"tasks": [...]

Incorrect:
"tasks": "[...]"
"""


def orchestrator_node(state: State) -> dict:
    response = planner_llm.invoke([
        SystemMessage(content=ORCH_SYSTEM),
        HumanMessage(content=(
            f"Topic: {state['topic']}\n"
            f"Mode: {state['mode']}\n"
            f"Evidence: {state['evidence']}"
        )),
    ])

    raw = response.content.strip()

    # Strip markdown fences
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    # ── NEW: Auto-repair broken JSON ──
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        # Try fixing with json_repair library
        try:
            from json_repair import repair_json
            data = json.loads(repair_json(raw))
        except Exception:
            raise ValueError(f"Planner returned invalid JSON:\n\n{raw}")

    try:
        plan = Plan.model_validate(data)
    except ValidationError as e:
        print(response.content)
        raise e

    return {"plan": plan}

# ─────────────────────────────────────────────
# Fan-Out
# ─────────────────────────────────────────────

def fanout(state: State):
    return [
        Send(
            "worker",
            {
                "task": task,
                "topic": state["topic"],
                "plan": state["plan"],
                "mode": state["mode"],                              # ← was missing before
                "evidence": [e.model_dump() for e in state.get("evidence", [])],
            },
        )
        for task in state["plan"].tasks
    ]


# ─────────────────────────────────────────────
# Node: Worker (Section Writer)
# ─────────────────────────────────────────────

def worker(payload: dict) -> dict:
    task: Task = payload["task"]
    topic: str = payload["topic"]
    plan: Plan = payload["plan"]
    mode: str = payload.get("mode", "closed_book")
    evidence: List[EvidenceItem] = [EvidenceItem(**e) for e in payload.get("evidence", [])]

    bullets_text = "\n- " + "\n- ".join(task.bullets)

    evidence_text = ""
    if evidence:
        evidence_text = "\n".join(
            f"- {e.title} | {e.url} | {e.published_at or 'date:unknown'}".strip()
            for e in evidence[:8]
        )

    section_md = writer_llm.invoke(
        [
            SystemMessage(
                content=(
                    "You are a senior technical writer. "
                    "Write ONE Markdown section for a technical blog.\n\n"
                    "Requirements:\n"
                    "- Cover the Goal and all Bullets in order.\n"
                    "- Stay within ±15% of the target word count.\n"
                    "- Output ONLY the section Markdown.\n\n"
                    "Writing Guidelines:\n"
                    "- Be technically accurate and practical.\n"
                    "- Use precise developer terminology.\n"
                    "- Include a short code example, checklist, or example when appropriate.\n"
                    "- Briefly mention trade-offs or edge cases if relevant.\n"
                    "- Explain why recommended practices matter.\n\n"
                    "Formatting:\n"
                    "- Start with '## Section Title'.\n"
                    "- Use short paragraphs, bullet lists, and fenced code blocks when useful.\n"
                    "- Avoid fluff and marketing language."
                )
            ),
            HumanMessage(
                content=(
                    f"Blog title: {plan.blog_title}\n"
                    f"Topic: {topic}\n"
                    f"Mode: {mode}\n"
                    f"Section title: {task.title}\n"
                    f"Goal: {task.goal}\n"
                    f"Target words: {task.target_words}\n"
                    f"Bullets:{bullets_text}\n\n"
                    f"Evidence (ONLY use these URLs when citing):\n{evidence_text}\n"
                )
            ),
        ]
    ).content.strip()

    return {"sections": [(task.id, section_md)]}


# ─────────────────────────────────────────────
# Node: Reducer (Fan-In)
# ─────────────────────────────────────────────

def reducer(state: State) -> dict:
    plan = state["plan"]

    ordered = sorted(state["sections"], key=lambda x: x[0])

    parts = [f"# {plan.blog_title}"]

    for _, section_md in ordered:
        parts.append(section_md)

    final_md = "\n\n".join(parts) + "\n"

    safe_title = "".join(
        c if c.isalnum() or c in (" ", "_", "-") else "" for c in plan.blog_title
    )
    filename = safe_title.strip().lower().replace(" ", "_") + ".md"

    output_dir = Path("blogs")
    output_dir.mkdir(exist_ok=True)

    doc = Document()

    doc.add_heading(plan.blog_title, level=1)

    for line in final_md.splitlines():

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)

        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)

        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")

        elif line.strip():
            doc.add_paragraph(line)

    safe_title = "".join(
        c if c.isalnum() or c in (" ", "_", "-") else ""
        for c in plan.blog_title
    )

    filename = safe_title.strip().lower().replace(" ", "_") + ".docx"

    output_dir = Path("blogs")
    output_dir.mkdir(exist_ok=True)

    docx_path = output_dir / filename
    doc.save(docx_path)

    return {
        "final": final_md,      # Keep this only for preview in Streamlit
        "docx_path": str(docx_path)
    }


# ─────────────────────────────────────────────
# Build & Compile Graph
# ─────────────────────────────────────────────

g = StateGraph(State)

g.add_node("router", router_node)
g.add_node("research", research_node)
g.add_node("orchestrator", orchestrator_node)
g.add_node("worker", worker)
g.add_node("reducer", reducer)

g.add_edge(START, "router")
g.add_conditional_edges(
    "router",
    route_next,
    {"research": "research", "orchestrator": "orchestrator"},
)
g.add_edge("research", "orchestrator")
g.add_conditional_edges("orchestrator", fanout, ["worker"])
g.add_edge("worker", "reducer")

app = g.compile()


# ─────────────────────────────────────────────
# Entry Point
# ─────────────────────────────────────────────

def run(topic: str) -> dict:
    out = app.invoke({
        "topic": topic,
        "mode": "",
        "needs_research": False,
        "queries": [],
        "evidence": [],
        "plan": None,
        "sections": [],
        "final": "",
        "docx_path": "",
    })
    return out

